#!/usr/bin/env python3
"""Generate a tailored resume as a professionally formatted DOCX file.

Usage:
    python generate_docx.py --original resume.docx --output tailored.docx --content tailored.json
    python generate_docx.py --template builtin --output tailored.docx --content tailored.json
    python generate_docx.py --original resume.docx --output tailored.docx --inline

The --inline mode reads rewritten content from stdin as a structured format.
Without --inline, provide a JSON file via --content with the rewritten sections.
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Formatting constants ──────────────────────────────────────────
FONT_ZH = "微软雅黑"
FONT_EN = "Calibri"
FONT_SIZE_NAME = Pt(22)
FONT_SIZE_SECTION = Pt(13)
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
COLOR_PRIMARY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_SECONDARY = RGBColor(0x2B, 0x5F, 0x8A)
COLOR_LIGHT_GRAY = RGBColor(0xE0, 0xE0, 0xE0)
COLOR_DIVIDER = RGBColor(0xCC, 0xCC, 0xCC)


def set_font(run, name=FONT_ZH, size=FONT_SIZE_BODY, color=COLOR_PRIMARY, bold=False):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    # Ensure CJK font rendering
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_divider(doc):
    """Add a thin horizontal divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_header(doc, title):
    """Add a section header with colored text and divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=FONT_SIZE_SECTION, color=COLOR_SECONDARY, bold=True)
    add_divider(doc)


def add_body_line(doc, text, bold_prefix=None, indent=False):
    """Add a body text line, optionally with a bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        set_font(run_b, bold=True)
    run = p.add_run(text)
    set_font(run)


def add_bullet(doc, text):
    """Add a bullet point line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    run = p.add_run("• " + text)
    set_font(run)


def build_header(doc, info):
    """Build the name + contact info header."""
    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(4)
    run = name_p.add_run(info.get("name", ""))
    set_font(run, size=FONT_SIZE_NAME, bold=True)

    # Contact line
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2)
    contact_items = []
    if info.get("email"):
        contact_items.append(f"📧 {info['email']}")
    if info.get("phone"):
        contact_items.append(f"📱 {info['phone']}")
    if info.get("location"):
        contact_items.append(f"📍 {info['location']}")
    if info.get("age"):
        contact_items.append(f"🎂 {info['age']}岁")
    if info.get("linkedin"):
        contact_items.append(f"🔗 {info['linkedin']}")
    run = contact_p.add_run("  |  ".join(contact_items))
    set_font(run, size=FONT_SIZE_SMALL, color=RGBColor(0x66, 0x66, 0x66))

    # Target role
    if info.get("target_role"):
        target_p = doc.add_paragraph()
        target_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        target_p.paragraph_format.space_after = Pt(6)
        run = target_p.add_run(f"求职意向：{info['target_role']}")
        set_font(run, size=FONT_SIZE_BODY, color=COLOR_SECONDARY)

    add_divider(doc)


def build_summary(doc, summary):
    """Build the professional summary section."""
    if not summary:
        return
    add_section_header(doc, "个人优势")
    add_body_line(doc, summary)


def build_skills(doc, skills):
    """Build the skills section as a compact table."""
    if not skills:
        return
    add_section_header(doc, "技能")

    # Determine columns: if skills is a list of strings, use 2 columns; if dict, use key-value
    if isinstance(skills, list):
        # Simple skill list — arrange in a 2-column table
        n = len(skills)
        mid = (n + 1) // 2
        col1 = skills[:mid]
        col2 = skills[mid:]
        table = doc.add_table(rows=max(len(col1), len(col2)), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, skill in enumerate(col1):
            cell = table.cell(i, 0)
            cell.text = ""
            run = cell.paragraphs[0].add_run(f"• {skill}")
            set_font(run)
        for i, skill in enumerate(col2):
            cell = table.cell(i, 1)
            cell.text = ""
            run = cell.paragraphs[0].add_run(f"• {skill}")
            set_font(run)
        # Remove table borders
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
    elif isinstance(skills, dict):
        # Category → skills mapping
        for category, skill_list in skills.items():
            text = f"{category}：{', '.join(skill_list) if isinstance(skill_list, list) else skill_list}"
            add_body_line(doc, text, bold_prefix=None)


def build_experience(doc, experiences):
    """Build the experience section."""
    if not experiences:
        return
    add_section_header(doc, "工作/实习经历")

    for exp in experiences:
        # Title line: Company — Role · Date
        title_parts = []
        if exp.get("company"):
            title_parts.append(exp["company"])
        if exp.get("role"):
            title_parts.append(exp["role"])
        if exp.get("date"):
            title_parts.append(exp["date"])
        title = " — ".join(title_parts[:-1]) if len(title_parts) > 1 else title_parts[0]
        if len(title_parts) > 1:
            title += f" · {title_parts[-1]}"
        elif len(title_parts) > 0:
            pass  # title already set

        # Actually, let me redo this
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(6)
        title_p.paragraph_format.space_after = Pt(2)
        run = title_p.add_run(f"{exp.get('company', '')} — {exp.get('role', '')} · {exp.get('date', '')}")
        set_font(run, bold=True)

        for bullet in exp.get("bullets", []):
            add_bullet(doc, bullet)


def build_projects(doc, projects):
    """Build the projects section."""
    if not projects:
        return
    add_section_header(doc, "项目经历")

    for proj in projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        title = proj.get("name", "")
        if proj.get("role"):
            title += f" — {proj['role']}"
        if proj.get("date"):
            title += f" · {proj['date']}"
        run = p.add_run(title)
        set_font(run, bold=True)

        for bullet in proj.get("bullets", []):
            add_bullet(doc, bullet)


def build_education(doc, education):
    """Build the education section."""
    if not education:
        return
    add_section_header(doc, "教育背景")

    if isinstance(education, dict):
        # Single education entry
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        text = f"{education.get('school', '')}  |  {education.get('degree', '')} · {education.get('major', '')}  |  {education.get('date', '')}"
        run = p.add_run(text)
        set_font(run, bold=True)
        for detail in education.get("details", []):
            add_body_line(doc, detail, indent=True)
    elif isinstance(education, list):
        for edu in education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            text = f"{edu.get('school', '')}  |  {edu.get('degree', '')} · {edu.get('major', '')}  |  {edu.get('date', '')}"
            run = p.add_run(text)
            set_font(run, bold=True)
            for detail in edu.get("details", []):
                add_body_line(doc, detail, indent=True)


def build_awards(doc, awards):
    """Build an awards/honors section."""
    if not awards:
        return
    add_section_header(doc, "获奖与荣誉")
    for award in awards:
        add_bullet(doc, award)


def build_from_content(doc, content):
    """Build full resume from structured content dict."""
    build_header(doc, content.get("header", {}))
    build_summary(doc, content.get("summary", ""))
    build_skills(doc, content.get("skills", {}))
    build_experience(doc, content.get("experience", []))
    build_projects(doc, content.get("projects", []))
    build_education(doc, content.get("education", {}))
    build_awards(doc, content.get("awards", []))


def load_or_create_content(content_arg, inline=False):
    """Load content from JSON file, stdin, or use inline structured data."""
    if inline:
        # Read from stdin
        raw = sys.stdin.read()
        return json.loads(raw)
    elif content_arg and Path(content_arg).exists():
        with open(content_arg, "r", encoding="utf-8") as f:
            return json.load(f)
    else:
        # content_arg might be a JSON string
        try:
            return json.loads(content_arg)
        except (json.JSONDecodeError, TypeError):
            raise ValueError(
                "Content must be a valid JSON file path or JSON string. "
                "Use --inline to pipe content via stdin."
            )


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate a tailored resume DOCX file"
    )
    parser.add_argument("--original", help="Path to original resume (for reference, not modified)")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument("--content", help="JSON file path or inline JSON with tailored content")
    parser.add_argument("--inline", action="store_true", help="Read content JSON from stdin")

    args = parser.parse_args()

    # Load content
    content = load_or_create_content(args.content, args.inline)

    # Create document
    doc = Document()

    # Set default page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Set default style
    style = doc.styles['Normal']
    style.font.name = FONT_ZH
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Build the resume
    build_from_content(doc, content)

    # Save
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Resume saved to: {output_path}")


if __name__ == "__main__":
    main()
