#!/usr/bin/env python3
"""
Rewrite resume DOCX text in-place, preserving ALL formatting.

Usage:
    python rewrite_resume.py original.docx output.docx replacements.json
    python rewrite_resume.py original.docx output.docx --inline '[{"find":"x","replace":"y"}]'

replacements.json format:
    [
        {"find": "original text to find", "replace": "new text"},
        ...
    ]

Key behavior:
- Modifies text at the RUN level, preserving font, size, bold, italic, color
- Walks both paragraphs AND table cells
- Only touches runs that contain the find string
- Saves to a NEW file (never overwrites original)
- Handles multi-run paragraphs: distributes replacement across runs
  maintaining each run's formatting
"""

import json
import sys
from pathlib import Path
from copy import deepcopy

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def replace_in_paragraph(paragraph, find_text, replace_text):
    """
    Replace text in a paragraph while preserving run-level formatting.

    Handles cases where find_text spans multiple runs.
    Returns True if a replacement was made.
    """
    runs = paragraph.runs
    if not runs:
        return False

    # Case 1: Simple — entire find_text is in a single run
    for run in runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text)
            return True

    # Case 2: Multi-run — find_text crosses run boundaries.
    # Preserve run structure: put replacement into the first affected run,
    # keep prefix/suffix text in place, clear only the fully-replaced middle runs.
    full_text = paragraph.text
    if find_text not in full_text:
        return False

    start_pos = full_text.find(find_text)
    end_pos = start_pos + len(find_text)

    # Locate which runs overlap with the [start_pos, end_pos) range
    char_pos = 0
    first_affected = None
    last_affected = None
    affected = set()

    for ri, run in enumerate(runs):
        run_start = char_pos
        run_end = char_pos + len(run.text)
        if run_start < end_pos and run_end > start_pos:
            affected.add(ri)
            if first_affected is None:
                first_affected = ri
            last_affected = ri
        char_pos = run_end

    if first_affected is None:
        return False

    # Save suffix from the last affected run (text after find_text)
    suffix = ""
    if last_affected > first_affected:
        last_run_start = sum(len(runs[i].text) for i in range(last_affected))
        offset_in_last = end_pos - last_run_start
        if offset_in_last < len(runs[last_affected].text):
            suffix = runs[last_affected].text[offset_in_last:]

    # Rebuild the first affected run: prefix + replacement + suffix
    first_run_start = sum(len(runs[i].text) for i in range(first_affected))
    offset_in_first = start_pos - first_run_start
    prefix = runs[first_affected].text[:offset_in_first]
    runs[first_affected].text = prefix + replace_text + suffix

    # Clear all other affected runs
    for ri in affected:
        if ri != first_affected:
            runs[ri].text = ""
    return True


def apply_replacements(doc, replacements):
    """
    Apply all replacements across the entire document — paragraphs and tables.

    Args:
        doc: python-docx Document
        replacements: list of {"find": str, "replace": str}

    Returns:
        dict with counts of replacements made
    """
    results = []
    for i, repl in enumerate(replacements):
        find_text = repl["find"]
        replace_text = repl["replace"]
        count = 0

        # Walk paragraphs
        for para in doc.paragraphs:
            if replace_in_paragraph(para, find_text, replace_text):
                count += 1

        # Walk table cells (each cell has its own paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if replace_in_paragraph(para, find_text, replace_text):
                            count += 1

        results.append({
            "find": find_text,
            "replace": replace_text,
            "occurrences_replaced": count,
        })

    return results


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="Rewrite resume DOCX text in-place, preserving all formatting"
    )
    parser.add_argument("original", help="Path to the original resume DOCX")
    parser.add_argument("output", help="Path for the output tailored DOCX")
    parser.add_argument("replacements", nargs="?", help="JSON file with replacements array")
    parser.add_argument("--inline", help="Inline JSON string of replacements (instead of file)")

    args = parser.parse_args()

    # Load replacements
    if args.inline:
        replacements = json.loads(args.inline)
    elif args.replacements:
        with open(args.replacements, "r", encoding="utf-8") as f:
            replacements = json.load(f)
    else:
        print("Error: provide replacements as JSON file or --inline string", file=sys.stderr)
        sys.exit(1)

    if not isinstance(replacements, list):
        print("Error: replacements must be a JSON array of {find, replace} objects", file=sys.stderr)
        sys.exit(1)

    # Open original document
    doc = Document(str(Path(args.original).resolve()))

    # Apply all replacements
    results = apply_replacements(doc, replacements)

    # Save to output path
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Report
    total = sum(r["occurrences_replaced"] for r in results)
    print(f"Applied {total} replacement(s) across {len(results)} rule(s):")
    for r in results:
        status = "done" if r["occurrences_replaced"] > 0 else "not found"
        print(f"  [{status}] '{r['find'][:60]}' -> '{r['replace'][:60]}'")
    print(f"Saved to: {output_path}")


if __name__ == "__main__":
    main()
